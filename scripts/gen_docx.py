#!/usr/bin/env python3
"""Parametric DOCX generator — reads company fit map + SHOOT package, outputs tailored DOCX."""

import sys
import os
import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from lxml import etree

ONEDRIVE = "/home/aryan/ABHIMANYU-2.0-output"
LINUX = "/home/aryan/opencode_test/ABHIMANYU-2.0"

NAME = "Aman Kumar"
PHONE = "+1 236-885-2285"
EMAIL = "amankumar7111@outlook.com"
LINKEDIN = "linkedin.com/in/aman1776"
LOCATION = "Vancouver, BC"

# HARD RULE: Companies with specific resume/cover letter content in generate().
# Adding a company to CONFIG without adding its content branch is FORBIDDEN.
# The script will FAIL if a company is in CONFIG but has no content branch AND no SHOOT package found.
_COMPANIES_WITH_CONTENT = {
    "Methanex", "Hiive", "Providence_Healthcare",
    "DoorDash_Canada", "UBC", "Practice_Better",
    "BWZ", "KPMG", "Microsoft",     "Indeed",
    "Accenture",
    "Brex",
    "Seaspan", "Clio",
}


def find_shoot_package(company):
    """Search date folders for SHOOT package for this company."""
    base_dirs = [
        LINUX,
        "/home/aryan/opencode_test",
    ]
    for base in base_dirs:
        for root, dirs, files in os.walk(base):
            for f in files:
                if f.endswith(".md") and company.lower() in f.lower() and ("shoot" in f.lower() or "01_" in f.lower() or f.startswith("01_")):
                    fp = os.path.join(root, f)
                    # Verify it's actually a SHOOT package for this company
                    try:
                        with open(fp) as fh:
                            content = fh.read(500)
                            if f"# SHOOT PACKAGE" in content or company in content:
                                return fp
                    except:
                        pass
    # Also check 01_CompanyName.md pattern in date folders
    import glob as gb
    for pat in [f"20*/WAVE_*/01_{company}.md", f"20*/WAVE_*/{company}.md"]:
        matches = gb.glob(os.path.join("/home/aryan/opencode_test/", pat))
        if matches:
            return matches[-1]  # most recent
    return None


def count_docx_content(doc):
    """Count meaningful content paragraphs in docx."""
    count = 0
    for p in doc.paragraphs:
        text = p.text.strip()
        if text and len(text) > 20:  # meaningful line
            count += 1
    return count

# Company-specific configs
CONFIG = {
    "Indeed": {
        "font": "Liberation Sans",
        "size": Pt(10),
        "header_size": Pt(12),
        "margins": Inches(0.75),
        "pages": 1,
        "ats_notes": "Indeed internal ATS — DOCX preferred, Liberation Sans, 0.75in margins"
    },
    "Methanex": {
        "font": "Calibri",
        "size": Pt(11),
        "header_size": Pt(13),
        "margins": Inches(0.75),
        "pages": 2,
        "ats_notes": "Methanex career portal — DOCX, Calibri 11pt, 2-page for Director level"
    },
    "Deloitte": {
        "font": "Calibri",
        "size": Pt(10),
        "header_size": Pt(12),
        "margins": Inches(0.75),
        "pages": 2,
        "ats_notes": "Deloitte Workday ATS — DOCX, Calibri 10pt"
    },
    "Hiive": {
        "font": "Calibri",
        "size": Pt(11),
        "header_size": Pt(13),
        "margins": Inches(0.75),
        "pages": 2,
        "ats_notes": "Hiive Ashby ATS — DOCX, Calibri 11pt"
    },
    "Providence_Healthcare": {
        "font": "Calibri",
        "size": Pt(11),
        "header_size": Pt(13),
        "margins": Inches(0.75),
        "pages": 2,
        "ats_notes": "Providence Oracle Cloud ATS — DOCX, Calibri 11pt"
    },
    "DoorDash_Canada": {
        "font": "Liberation Sans",
        "size": Pt(10),
        "header_size": Pt(12),
        "margins": Inches(0.75),
        "pages": 1,
        "ats_notes": "DoorDash Greenhouse ATS — DOCX, Liberation Sans 10pt, 0.75in margins"
    },
    "UBC": {
        "font": "Liberation Sans",
        "size": Pt(10),
        "header_size": Pt(12),
        "margins": Inches(0.75),
        "pages": 1,
        "ats_notes": "UBC / Indeed ATS — DOCX, Liberation Sans 10pt, 0.75in margins, public sector format"
    },
    "Practice_Better": {
        "font": "Liberation Sans",
        "size": Pt(10),
        "header_size": Pt(12),
        "margins": Inches(0.75),
        "pages": 2,
        "ats_notes": "Practice Better Greenhouse ATS — DOCX, Liberation Sans 10pt, 0.75in margins"
    },
    "RAM_Consulting": {
        "font": "Liberation Sans",
        "size": Pt(10),
        "header_size": Pt(12),
        "margins": Inches(0.75),
        "pages": 2,
        "ats_notes": "RAM Consulting DayforceHCM ATS — DOCX, Liberation Sans 10pt, 0.75in margins, project delivery focus"
    },
    "BWZ": {
        "font": "Calibri",
        "size": Pt(10),
        "header_size": Pt(12),
        "margins": Inches(0.75),
        "pages": 2,
        "ats_notes": "BWZ Lever ATS — DOCX, Calibri 10pt, 0.75in margins, strategy & operations focus"
    },
    "KPMG": {
        "font": "Liberation Sans",
        "size": Pt(10),
        "header_size": Pt(12),
        "margins": Inches(0.75),
        "pages": 2,
        "ats_notes": "KPMG ICIMS ATS — DOCX, Liberation Sans 10pt, 0.75in margins, Director program delivery"
    },
    "Microsoft": {
        "font": "Segoe UI",
        "size": Pt(10),
        "header_size": Pt(12),
        "margins": Inches(0.75),
        "pages": 2,
        "ats_notes": "Microsoft Workday ATS — DOCX, Segoe UI 10pt, 0.75in margins, Sr TPM Security"
    },
    "Accenture": {
        "font": "Calibri",
        "size": Pt(10),
        "header_size": Pt(12),
        "margins": Inches(0.75),
        "pages": 2,
        "ats_notes": "Accenture Workday — DOCX, Calibri 10pt"
    },
    "Brex": {
        "font": "Calibri",
        "size": Pt(10),
        "header_size": Pt(12),
        "margins": Inches(0.75),
        "pages": 2,
        "ats_notes": "Brex Greenhouse ATS — DOCX, Calibri 10pt, 0.75in margins, BizOps Sr Mgr (Technical)"
    },
    "MASTER": {
        "font": "Calibri",
        "size": Pt(11),
        "header_size": Pt(13),
        "margins": Inches(0.75),
        "pages": 2,
        "ats_notes": "Master resume for executive search firms — Calibri 11pt, 2-page"
    }
}

def set_margins(doc, margin):
    for section in doc.sections:
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin

def add_contact(doc, config):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(NAME)
    run.font.name = config["font"]
    run.font.size = Pt(16)
    run.bold = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_plain_run(p2, f"{PHONE}  |  ", config["font"], Pt(9), color="505050")
    add_hyperlink_contact(p2, EMAIL, f"mailto:{EMAIL}", config["font"], Pt(9))
    add_plain_run(p2, "  |  ", config["font"], Pt(9), color="505050")
    add_hyperlink_contact(p2, "LinkedIn", f"https://{LINKEDIN}", config["font"], Pt(9))
    add_plain_run(p2, f"  |  {LOCATION}", config["font"], Pt(9), color="505050")

def add_section_header(doc, text, config):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    run.font.name = config["font"]
    run.font.size = config["header_size"]
    run.bold = True
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '4',
        qn('w:space'): '1', qn('w:color'): '000000',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_body(doc, text, config, bold=False, italic=False, size=None, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name = config["font"]
    run.font.size = size or config["size"]
    run.bold = bold
    run.italic = italic

def add_bullet(doc, text, config, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    if bold_prefix:
        run_b = p.add_run(f"• {bold_prefix}")
        run_b.font.name = config["font"]
        run_b.font.size = config["size"]
        run_b.bold = True
        run = p.add_run(text)
        run.font.name = config["font"]
        run.font.size = config["size"]
    else:
        run = p.add_run(f"• {text}")
        run.font.name = config["font"]
        run.font.size = config["size"]

def add_hyperlink_contact(p, label, url, font_name, font_size):
    part = p.part
    r_id = part.relate_to(url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True)
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    ns_r = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    hyperlink_xml = (
        f'<w:hyperlink xmlns:w="{ns_w}" xmlns:r="{ns_r}" '
        f'r:id="{r_id}" w:history="1">'
        f'<w:r><w:rPr>'
        f'<w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}"/>'
        f'<w:sz w:val="{int(font_size.pt * 2)}"/>'
        f'<w:color w:val="0563C1"/>'
        f'<w:u w:val="single"/>'
        f'</w:rPr>'
        f'<w:t xml:space="preserve">{label}</w:t>'
        f'</w:r></w:hyperlink>'
    )
    hyperlink_elem = parse_xml(hyperlink_xml)
    p._p.append(hyperlink_elem)

def add_plain_run(p, text, font_name, font_size, color=None):
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    color_attr = f'<w:color w:val="{color}"/>' if color else ''
    run_xml = (
        f'<w:r xmlns:w="{ns_w}">'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}"/>'
        f'<w:sz w:val="{int(font_size.pt * 2)}"/>'
        f'{color_attr}'
        f'</w:rPr>'
        f'<w:t xml:space="preserve">{text}</w:t>'
        f'</w:r>'
    )
    run_elem = parse_xml(run_xml)
    p._p.append(run_elem)

def add_signature(doc, config):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    add_plain_run(p, "Best regards,", config["font"], config["size"])
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.space_before = Pt(0)
    add_plain_run(p2, NAME, config["font"], config["size"])
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(0)
    p3.paragraph_format.space_before = Pt(0)
    add_plain_run(p3, PHONE, config["font"], config["size"], color="505050")
    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(0)
    p4.paragraph_format.space_before = Pt(0)
    add_hyperlink_contact(p4, EMAIL, f"mailto:{EMAIL}", config["font"], config["size"])

def get_date(company):
    DATES = {
        "Indeed": "2026-07-05",
        "Methanex": "2026-06-21",
        "Deloitte": "2026-06-19",
        "Hiive": "2026-06-22",
        "Providence_Healthcare": "2026-06-22",
        "DoorDash_Canada": "2026-06-24",
        "UBC": "2026-06-25",
        "Practice_Better": "2026-06-25",
        "RAM_Consulting": "2026-06-25",
        "BWZ": "2026-06-25",
        "KPMG": "2026-06-25",
        "Microsoft": "2026-06-25",
        "Accenture": "2026-07-07",
        "Brex": "2026-07-15",
        "Seaspan": "2026-07-14",
    }
    return DATES.get(company, "2026-06-22")

def generate(company):
    # DELEGATION: MASTER now handled by resume_docx_builder.py
    if company == "MASTER":
        print("MASTER: delegating to resume_docx_builder.py")
        script_dir = os.path.dirname(os.path.abspath(__file__))
        source_path = os.path.join(os.path.dirname(script_dir), "Master_Resume_ATS_CLEAN.md")
        subprocess.run([sys.executable,
            os.path.join(script_dir, "resume_docx_builder.py"),
            "--source", source_path,
            "--company", "MASTER",
            "--role", "Master_Resume"
        ])
        return

    config = CONFIG.get(company, CONFIG["Methanex"])
    date_str = get_date(company)
    folder = f"{ONEDRIVE}/{date_str}/{company}"
    lfolder = f"{LINUX}/{date_str}/{company}"
    os.makedirs(folder, exist_ok=True)
    os.makedirs(lfolder, exist_ok=True)

    # --- SAFETY CHECK: Ensure company has content or SHOOT package ---
    shoot_pkg = find_shoot_package(company)
    has_branch = company in _COMPANIES_WITH_CONTENT
    if not has_branch and not shoot_pkg:
        print(f"ERROR: '{company}' has no content branch in generate() and no SHOOT package found.")
        print(f"  → Add to _COMPANIES_WITH_CONTENT and add an 'elif company == \"{company}\":' block in generate()")
        print(f"  → Or create a SHOOT package (01_{company}.md) in a date folder first.")
        sys.exit(1)
    if not has_branch and shoot_pkg:
        print(f"WARNING: '{company}' has no content branch but SHOOT package found at {shoot_pkg}")
        print(f"  → Using generic content. Add a content branch in generate() for full tailoring.")
        # Continue anyway — SHOOT package exists so we have real resume text available

    # --- RESUME ---
    doc = Document()
    set_margins(doc, config["margins"])
    add_contact(doc, config)

    add_section_header(doc, "Professional Summary", config)
    if company == "Methanex":
        add_body(doc,
            "Strategy + operations executive. Scaled 3→70, 32 locations, $4M ARR, $17M exit. "
            "Built strategic planning, financial models (multi-scenario P&L, valuation, capital allocation), "
            "and board-level reporting from zero. Boardroom-fluent. Builder-grounded.", config)
    elif company == "Hiive":
        add_body(doc,
            "Operations strategist + systems builder. Scaled 3→70, 32 locations, $4M ARR, $17M exit. "
            "Designed operational infrastructure, automated revenue lifecycle, and built data systems "
            "from zero. Eliminated bottlenecks. Never added complexity.", config)
    elif company == "Providence_Healthcare":
        add_body(doc,
            "Healthcare operations leader. Scaled 3→70, 32 locations, $4M ARR, $17M exit. "
            "Directed multi-site clinical ops, interdisciplinary teams, and P&L across 12 departments. "
            "Geriatrics-specialized. Quality improvement. $4M budget ownership.", config)
    elif company == "DoorDash_Canada":
        add_body(doc,
            "Operations executive. Scaled 3→70, 32 locations, $4M ARR, $17M exit. "
            "Designed marketplace incentive architecture across 32 markets — balancing worker earnings "
            "with cost efficiency. Led technology transformation. Full P&L ownership.", config)
    elif company == "Clio":
        # Executive‑grade professional summary for a Fortune‑500‑level ops leader
        add_body(doc,
            "Seasoned technology‑operations executive with a track record of scaling high‑growth organizations from 3 to 70+ employees, “zero‑to‑one” product launches, and $17 M exits. "
            "Deep expertise in SaaS, cloud‑native infrastructures, data‑pipeline architecture, and cross‑functional team leadership across 32 sites. "
            "Proven ability to drive P&L ownership, operational excellence, and strategic transformation in Fortune‑500 environments.", config)
        # Core competencies (compact bullet list)
        add_section_header(doc, "Core Competencies", config)
        add_bullet(doc, "Strategic Operations & Scaling", config, bold_prefix="• ")
        add_bullet(doc, "SaaS & Cloud Architecture", config, bold_prefix="• ")
        add_bullet(doc, "Data Engineering & Pipeline Design", config, bold_prefix="• ")
        add_bullet(doc, "P&L Management & Financial Modeling", config, bold_prefix="• ")
        add_bullet(doc, "Cross‑Functional Team Leadership", config, bold_prefix="• ")
        add_bullet(doc, "M&A Integration & Program Governance", config, bold_prefix="• ")
    elif company == "Practice_Better":
        add_body(doc,
            "Revenue operations architect. Scaled 3→70, 32 locations, $4M ARR, $17M exit. "
            "Built complete RevOps stack from zero — EHR, billing, RCM, analytics, forecasting. "
            "SaaS metrics (MRR, ARR, churn, CAC, LTV). End-to-end pipeline ownership.", config)
    elif company == "BWZ":
        add_body(doc,
            "Zero-to-one operator. Scaled 3→70, 32 locations, $4M ARR, $17M exit. "
            "Built initiatives from idea through business case to market launch. "
            "Program discipline without bureaucracy. AI-augmented. Founder DNA.", config)
    elif company == "KPMG":
        add_body(doc,
            "Operations + program delivery executive. Scaled 3→70, 32 locations, $4M ARR, $17M exit. "
            "Builder's execution + consultant's strategic framing. Program governance, digital "
            "transformation, P&L ownership. MBA. Builds, then advises.", config)
    elif company == "Indeed":
        add_body(doc,
            "Operations + integration leader. Scaled 3→70, 32 locations, $4M ARR, $17M exit. "
            "Led M&A integration, cross-functional programs, and technology transformations. "
            "Built integration playbooks from scratch. P&L owner. Executive communicator.", config)
    elif company == "Accenture":
        add_body(doc,
            "Strategy + operations executive. Scaled 3→70 across 32 locations in 4 US states, "
            "$4M ARR, $17M exit. Built centralized operations backbone — tech, team, process — "
            "from zero. P&L owner through hypergrowth. Board-level communicator.", config)
    elif company == "Brex":
        add_body(doc,
            "BUILT FROM ZERO: 3→70 FTE · $300K→$4M ARR · 1→32 locations · $17M exit. "
            "Operations builder — designed every system (EHR, VOIP, automation, Python pipelines), "
            "led AI-driven transformation, directed the acquisition. Board-communicator. Hands-on builder.", config)
    elif company == "Seaspan":
        add_body(doc,
            "Change management practitioner + operations builder. Scaled 3→70, 32 locations, "
            "$4M ARR, $17M exit. Led organizational change through 5 acquisitions — building "
            "adoption frameworks, stakeholder systems, and integration playbooks that retained "
            "100% key talent. PROSCI-aligned. ADKAR-practiced.", config)
    elif company == "Microsoft":
        add_body(doc,
            "Technical program management leader. Scaled 3→70, 32 locations, $4M ARR, $17M exit. "
            "Defined operational readiness frameworks, led technology transformations, managed $4M P&L. "
            "Executive communication. AI adoption at scale.", config)
    else:
        add_body(doc,
            "Operations executive. Scaled 3→70, 32 locations, $4M ARR, $17M exit. "
            "Built multi-site infrastructure from scratch. Strategic + execution.", config)

    add_section_header(doc, "Core Competencies", config)
    if company == "Methanex":
        add_body(doc,
            "Corporate Strategy · Financial Modelling & Valuation · M&A Execution & Integration · "
            "Board & Executive Communication · Capital Allocation · OKR & Performance Systems · "
            "Cross-Functional Leadership · Strategic Growth Initiatives", config, size=Pt(9.5))
    elif company == "Hiive":
        add_body(doc,
            "Revenue Operations · Systems Architecture & Automation · "
            "Bottleneck Analysis · Workflow Optimization · "
            "Data Integrity & Reporting · Cross-Functional Execution · "
            "M&A & Strategic Projects · Operational Infrastructure Design", config, size=Pt(9.5))
    elif company == "Providence_Healthcare":
        add_body(doc,
            "Multi-Site Healthcare Ops · Clinical Operations Leadership · Quality Improvement · "
            "Financial Management & Budgeting ($4M) · Interdisciplinary Team Leadership · "
            "Change Management · Regulatory Compliance · Geriatrics & Senior Care", config, size=Pt(9.5))
    elif company == "DoorDash_Canada":
        add_body(doc,
            "Marketplace Operations · Pay & Incentive Design · P&L Management · "
            "Cross-Functional Leadership · Behavioral Economics · Data-Driven Strategy · "
            "0-to-1 Scaling · Operational Excellence", config, size=Pt(9.5))
    elif company == "Practice_Better":
        add_body(doc,
            "Revenue Operations · SaaS Metrics (MRR/ARR/Churn/CAC/LTV) · P&L Management · "
            "Revenue Cycle Management · GTM Alignment · Pipeline & Forecasting · "
            "Workflow Automation · KPI Dashboard Design", config, size=Pt(9.5))
    elif company == "BWZ":
        add_body(doc,
            "Zero-to-One Initiative Build · Program & Project Discipline · Business Case Development · "
            "Cross-Functional Leadership · Financial Modeling (Revenue/Cost/Margin) · P&L Management · "
            "AI-Augmented Workflows · Operational Infrastructure Design", config, size=Pt(9.5))
    elif company == "KPMG":
        add_body(doc,
            "Program Delivery & Governance · Digital Transformation · P&L Management · "
            "Cross-Functional Leadership · Strategic Planning & OKRs · M&A & Integration · "
            "Risk Management · Board-Level Reporting & Stakeholder Engagement", config, size=Pt(9.5))
    elif company == "Microsoft":
        add_body(doc,
            "Technical Program Management · Operational Readiness Frameworks · "
            "Cross-Functional Leadership · Technology Transformation · P&L Management · "
            "Executive Communication · M&A Integration · AI-Augmented Workflows", config, size=Pt(9.5))
    elif company == "Accenture":
        add_body(doc,
            "Strategic Transformation · Technology Enablement · Operational Resilience · "
            "Cross-Functional Leadership · M&A & Integration · P&L Management · "
            "Board-Level Communication · Business Case Development", config, size=Pt(9.5))
    elif company == "Brex":
        add_body(doc,
            "Systems Architecture · AI-Native Automation · Decision Infrastructure · "
            "Scalable Platforms · Workflow Design · Internal Tooling · "
            "Cross-Functional Leadership · Acquisition Integration", config, size=Pt(9.5))
    elif company == "Indeed":
        add_body(doc,
            "M&A Integration · Cross-Functional Program Management · Operational Infrastructure · "
            "Strategic Planning & OKRs · P&L Management · Board Reporting · Multi-Site Operations",
            config, size=Pt(9.5))
    elif company == "Seaspan":
        add_body(doc,
            "Change Management & Adoption · Organizational Transformation · Stakeholder Engagement · "
            "M&A Integration & Culture · Strategic Planning & OKRs · Cross-Functional Leadership · "
            "Process Improvement · Operational Infrastructure Design", config, size=Pt(9.5))
    add_section_header(doc, "Professional Experience", config)

    # SkyflyMD
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run("SkyflyMD")
    run.font.name = config["font"]
    run.font.size = config["size"]
    run.bold = True
    run2 = p.add_run("  |  Director of Operations  |  Phoenix, AZ / Vancouver, BC  |  2017 – 2025")
    run2.font.name = config["font"]
    run2.font.size = config["size"]

    if company == "Methanex":
        add_body(doc,
            "Built strategic planning, financial infrastructure, and operational systems from zero — "
            "from 3 people and no playbook to 70 FTEs, 32 locations, and a $17M exit.", config, italic=True, size=Pt(9.5))
        add_bullet(doc,
            " — built company-wide strategic planning: board-level strategy sessions, OKR cascades, "
            "quarterly reviews aligned across 5 clinic groups, 12 departments, 32 locations",
            config, bold_prefix="Strategic Planning & Execution")
        add_bullet(doc,
            " — multi-scenario P&L models, capital allocation frameworks, departmental budgets across "
            "12 departments — governed $4M in annual resource allocation",
            config, bold_prefix="Financial Modelling & Capital Allocation")
        add_bullet(doc,
            " — directed end-to-end $17M acquisition: 8 diligence workstreams, Day 1 readiness, "
            "90-day systems consolidation, 100% key talent retention",
            config, bold_prefix="M&A Execution & Integration")
        add_bullet(doc,
            " — built organizational infrastructure: hiring frameworks, training programs, quality "
            "standards, cross-border coordination for 70 FTEs across 5 clinic groups",
            config, bold_prefix="Organizational Leadership & Team Scaling")
        add_bullet(doc,
            " — presented board-ready financial reporting and strategic updates to executive leadership "
            "and investors — governed $4M annual budget with monthly variance tracking",
            config, bold_prefix="Board-Level Reporting & Governance")
    elif company == "Hiive":
        add_body(doc,
            "Designed operational infrastructure from zero — systems, processes, data pipelines — "
            "enabling 23x growth without adding complexity.", config, italic=True, size=Pt(9.5))
        add_bullet(doc,
            " — designed and implemented the entire operational tech stack (EHR, billing, scheduling, "
            "RCM, analytics) that scaled from 3→70 across 32 locations without adding complexity",
            config, bold_prefix="Systems Architecture & Automation")
        add_bullet(doc,
            " — eliminated revenue lifecycle bottlenecks: automated billing workflows, replaced manual "
            "reconciliation with real-time RCM, cut admin overhead by 40%+",
            config, bold_prefix="Revenue Operations & Optimization")
        add_bullet(doc,
            " — built KPI dashboards, board-level reporting, and data integrity systems from scratch — "
            "replaced manual spreadsheets with real-time analytics across all 32 locations",
            config, bold_prefix="Data Infrastructure & Reporting")
        add_bullet(doc,
            " — directed full-cycle $17M acquisition: 8 diligence workstreams, integration playbook, "
            "consolidated 8 systems within 90 days, 100% key talent retention",
            config, bold_prefix="M&A & Strategic Execution")
    elif company == "Providence_Healthcare":
        add_body(doc,
            "Directed clinical + operational leadership across 32 locations, 12 departments, "
            "and 5 clinic groups. Full P&L ownership. Geriatrics-specialized.", config, italic=True, size=Pt(9.5))
        add_bullet(doc,
            " — coordinated interdisciplinary teams across 5 clinic groups delivering patient-centered "
            "geriatric care — multi-facility coordination including senior care homes",
            config, bold_prefix="Multi-Site Healthcare Operations")
        add_bullet(doc,
            " — managed full P&L ownership for $4M healthcare org — budget planning, variance analysis, "
            "resource allocation, capital expenditure across 32 locations",
            config, bold_prefix="Financial Management & Resource Allocation")
        add_bullet(doc,
            " — drove quality improvement: automated workflows, standardized processes, data-driven "
            "performance management — reduced administrative overhead by 40%+",
            config, bold_prefix="Quality Improvement & Process Optimization")
        add_bullet(doc,
            " — directed end-to-end $17M acquisition: 8 diligence workstreams, Day 1 readiness, "
            "90-day systems consolidation, 100% key talent retention across 32 locations",
            config, bold_prefix="Change Management & Transformation")
        add_bullet(doc,
            " — managed geriatric practice operations: scheduling, billing (ICD coding, insurance claims), "
            "multi-site coordination across 5+ clinics and multiple senior care homes",
            config, bold_prefix="Geriatric & Senior Care Operations")
    elif company == "DoorDash_Canada":
        add_body(doc,
            "Built from zero: 3→70 people, 32 locations, $4M ARR, $17M exit. "
            "Owned P&L, designed incentive architecture, led technology transformation.", config, italic=True, size=Pt(9.5))
        add_bullet(doc,
            " — scaled 3→70 with infrastructure that supported 23x growth without adding complexity — "
            "hiring frameworks, training, cross-functional processes",
            config, bold_prefix="Operational Scaling & Infrastructure")
        add_bullet(doc,
            " — designed incentive compensation across 32 markets — tuned per location for labour cost, "
            "competitive pressure, and worker expectations while maintaining budget discipline",
            config, bold_prefix="Pay & Incentive Design (Marketplace)")
        add_bullet(doc,
            " — directed full-cycle $17M acquisition: 8 diligence workstreams, Day 1 readiness, "
            "90-day systems consolidation, 100% key talent retention",
            config, bold_prefix="M&A & Exit Execution")
        add_bullet(doc,
            " — managed $3M+ annual budget, multi-scenario forecasting, variance analysis, capital "
            "allocation across 12 departments — presented board-ready reporting to investors",
            config, bold_prefix="P&L Management & Financial Operations")
        add_bullet(doc,
            " — led transformation from paper to integrated digital platform: EHR, billing, scheduling, "
            "analytics — selected, deployed, and owned every system",
            config, bold_prefix="Technology Transformation")
        add_bullet(doc,
            " — implemented KPI dashboards across 32 locations, reducing reporting lag by 30%. "
            "Negotiated vendor contracts achieving 12% cost savings.",
            config, bold_prefix="KPI Dashboards & Vendor Management")
    elif company == "UBC":
        add_body(doc,
            "Built strategic governance, stakeholder systems, and financial infrastructure from zero — "
            "serving as the bridge between executive vision and operational execution.", config, italic=True, size=Pt(9.5))
        add_bullet(doc,
            " — built strategic planning framework across 5 annual cycles: board-level strategy sessions, "
            "OKR cascades, quarterly reviews — aligning 5 clinic groups, 12 departments, 32 locations",
            config, bold_prefix="Strategic Planning & Governance")
        add_bullet(doc,
            " — managed full P&L ownership for $4M ARR organization: department budgets, variance "
            "analysis, multi-scenario forecasting, board-ready financial reporting",
            config, bold_prefix="Budget & Financial Management")
        add_bullet(doc,
            " — directed end-to-end $17M acquisition: 8 diligence workstreams, Day 1 readiness, "
            "90-day systems consolidation across 32 locations, 100% key talent retention",
            config, bold_prefix="Acquisition & Integration Leadership")
        add_bullet(doc,
            " — built governance framework satisfying HIPAA and state licensing across 5 jurisdictions — "
            "internal accountability and external regulatory compliance",
            config, bold_prefix="Compliance & Regulatory Governance")
        add_bullet(doc,
            " — stakeholder engagement across 12 departments, 5 clinic groups with competing priorities: "
            "monthly check-in cadences, escalation frameworks, prioritization matrices",
            config, bold_prefix="Stakeholder Engagement & Alignment")
        add_bullet(doc,
            " — presented quarterly board-level reporting to executive leadership — P&L variance, "
            "strategic initiative tracking, resource allocation governing $4M annual budget",
            config, bold_prefix="Board Reporting & Executive Communication")
    elif company == "Practice_Better":
        add_body(doc,
            "Architected complete RevOps infrastructure from zero — systems, metrics, pipeline — "
            "scaling from $0→$4M ARR across 32 locations through a $17M exit.", config, italic=True, size=Pt(9.5))
        add_bullet(doc,
            " — built end-to-end RevOps stack from scratch: EHR, billing, scheduling, RCM, analytics — "
            "scaled 3→70 across 32 locations without adding complexity",
            config, bold_prefix="Revenue Operations Infrastructure")
        add_bullet(doc,
            " — architected SaaS metrics framework: MRR, ARR, churn, CAC, LTV — built KPI dashboards "
            "replacing manual spreadsheets, reduced reporting lag by 30%",
            config, bold_prefix="SaaS Metrics & Data Infrastructure")
        add_bullet(doc,
            " — eliminated revenue lifecycle bottlenecks: automated billing workflows, replaced manual "
            "reconciliation with real-time RCM, cut admin overhead by 40%+",
            config, bold_prefix="Pipeline Management & Optimization")
        add_bullet(doc,
            " — managed full P&L ownership for $4M ARR: budget planning, variance analysis, resource "
            "allocation, multi-scenario forecasting across 12 departments",
            config, bold_prefix="P&L Ownership & Forecasting")
        add_bullet(doc,
            " — directed full-cycle $17M acquisition: 8 diligence workstreams, integration playbook, "
            "Day 1/100 milestones, 8→1 systems consolidation, 100% talent retention",
            config, bold_prefix="M&A Revenue Integration")
        add_bullet(doc,
            " — coordinated GTM alignment across sales, ops, finance, clinical teams — reduced pipeline "
            "variance from 40% to under 10% through unified reporting",
            config, bold_prefix="GTM Alignment & Cross-Functional Leadership")
    elif company == "BWZ":
        add_body(doc,
            "Built from zero to 70 people, 32 locations, $4M ARR — then directed the $17M exit. "
            "Primary operator, systems architect, cross-functional integrator.", config, italic=True, size=Pt(9.5))
        add_bullet(doc,
            " — owned end-to-end zero-to-one initiatives: from idea through costing, pricing, GTM "
            "to launch — across 32 locations and 5 clinic groups",
            config, bold_prefix="Zero-to-One Initiative Build")
        add_bullet(doc,
            " — installed project discipline: governance frameworks with clear owners, documented "
            "decisions, cross-functional follow-through — without adding bureaucracy",
            config, bold_prefix="Program & Project Discipline")
        add_bullet(doc,
            " — acted as business-case clearinghouse: pressure-tested half-baked ideas into board-ready "
            "proposals with revenue projections, cost structures, break-even analysis",
            config, bold_prefix="Business Case Development")
        add_bullet(doc,
            " — built financial models for every new initiative: revenue projections, cost structures, "
            "sensitivity analysis — from napkin math to board-ready",
            config, bold_prefix="Financial Modeling & Analysis")
        add_bullet(doc,
            " — managed full P&L: budget planning, variance analysis, multi-scenario forecasting "
            "across 12 departments and 32 locations",
            config, bold_prefix="P&L Management & Forecasting")
        add_bullet(doc,
            " — led AI-driven automation: transformed billing, scheduling, reporting — cut admin "
            "overhead by 40%+ using AI tools daily",
            config, bold_prefix="AI-Augmented Operations")
    elif company == "KPMG":
        add_body(doc,
            "Scaled 3→70, 32 locations, $4M ARR, $17M exit. Builder who executed + strategist "
            "who designed the delivery framework.", config, italic=True, size=Pt(9.5))
        add_bullet(doc,
            " — directed end-to-end $17M acquisition: 8 diligence workstreams across finance, legal, "
            "ops — built integration playbook, Day 1/100 milestones, 8→1 systems, 100% talent retention",
            config, bold_prefix="Program Delivery & Governance")
        add_bullet(doc,
            " — led complete digital transformation: designed EHR, billing, scheduling, RCM, analytics "
            "platform across 32 locations — replaced paper with integrated tech stack",
            config, bold_prefix="Digital Transformation")
        add_bullet(doc,
            " — managed full P&L: department budgets, variance analysis, multi-scenario forecasting, "
            "capital allocation — board-ready reporting to executive leadership and investors",
            config, bold_prefix="P&L Management & Financial Reporting")
        add_bullet(doc,
            " — built strategic planning and governance framework: annual strategy cycles, OKR cascades, "
            "quarterly reviews — across 5 clinic groups, 12 departments, 5 regulatory jurisdictions",
            config, bold_prefix="Strategic Planning & Risk Management")
        add_bullet(doc,
            " — led cross-functional teams across 12 departments — managed 70 FTEs through hypergrowth, "
            "built hiring frameworks, training programs, quality standards",
            config, bold_prefix="Cross-Functional Leadership & Team Building")
    elif company == "Microsoft":
        add_body(doc,
            "Scaled 3→70, 32 locations, $4M ARR, $17M exit. Primary operator, systems architect, "
            "and program manager defining how work got done at scale.", config, italic=True, size=Pt(9.5))
        add_bullet(doc,
            " — defined operational readiness frameworks and governance processes across clinical, "
            "financial, compliance, and technology domains — enabling growth from 3→70 across 32 locations",
            config, bold_prefix="Operational Readiness Frameworks")
        add_bullet(doc,
            " — led cross-functional technology transformation: designed EHR, billing, scheduling, RCM "
            "platform — managed vendor selection, stakeholder alignment, rollout sequencing",
            config, bold_prefix="Technology Transformation Programs")
        add_bullet(doc,
            " — managed $4M ARR P&L, built board-level reporting dashboards, investor communications, "
            "executive strategy cadence — translated operational signals into $17M exit insights",
            config, bold_prefix="Program Metrics & Executive Communication")
        add_bullet(doc,
            " — led integration programs spanning M&A, technology, compliance, ops — aligned priorities "
            "across 7+ departments with governance frameworks and clear decision rights",
            config, bold_prefix="Cross-Organizational Program Leadership")
        add_bullet(doc,
            " — designed customer feedback loops, KPI frameworks, and dashboards that informed "
            "product roadmaps, service improvements, and strategic resource allocation",
            config, bold_prefix="Customer Feedback → Product Strategy")
    elif company == "Accenture":
        add_body(doc,
            "Built centralized operations backbone for a multi-site roll-up from scratch — "
            "3→70 FTEs, 32 locations, 4 US states, 5+ acquisitions, $17M exit.", config, italic=True, size=Pt(9.5))
        add_bullet(doc,
            " — directed end-to-end $17M acquisition: 8 diligence workstreams across finance, legal, "
            "ops, provider contracts — integration playbook with Day 1/100 milestones, 8→1 systems, "
            "100% talent retention across 32 locations",
            config, bold_prefix="Strategic Transformation & M&A")
        add_bullet(doc,
            " — designed and implemented complete tech stack (EHR, billing, scheduling, RCM, analytics) "
            "that scaled 3→70 without adding complexity — replaced paper with integrated digital systems",
            config, bold_prefix="Technology Enablement & Systems Architecture")
        add_bullet(doc,
            " — managed full P&L: multi-scenario financial models, department budgets, variance analysis, "
            "board-ready reporting — real-time visibility across 12 departments and 32 locations",
            config, bold_prefix="Value Creation & P&L Management")
        add_bullet(doc,
            " — built strategic planning from zero: annual strategy cycles, quarterly OKRs, board-level "
            "reporting — aligned 5 clinic groups, 12 departments, 32 locations across 5 annual cycles",
            config, bold_prefix="Performance Improvement & Strategic Planning")
        add_bullet(doc,
            " — led cross-functional team of 45-70 FTEs: recruitment, training, performance management, "
            "organizational design — turned a startup into a scalable enterprise",
            config, bold_prefix="Cross-Functional Leadership & Organizational Design")
    elif company == "Indeed":
        add_body(doc,
            "Operations Lead, Multi-Site Integration — managed integration across 5+ clinic groups, "
            "32 locations in AZ and TX. Ensured Day 1 readiness for every acquisition.", config, italic=True, size=Pt(9.5))
        add_bullet(doc,
            " — established integration governance: workstream cadences, milestone tracking, escalation "
            "paths across 5+ clinic groups spanning 32 locations — zero operational disruption on Day 1",
            config, bold_prefix="Multi-Site Integration Governance")
        add_bullet(doc,
            " — built standardized integration playbooks from scratch — documented every process for "
            "repeatable execution — new locations reached steady state within 30 days of close",
            config, bold_prefix="Integration Playbooks & Frameworks")
        add_bullet(doc,
            " — led cross-functional team of 45-60 across scheduling, clinical coordination, outreach, "
            "compliance, IT — aligned execution without direct authority",
            config, bold_prefix="Cross-Functional Team Leadership")
        add_bullet(doc,
            " — translated diligence insights into actionable plans with clear owners — tracked progress "
            "against synergy targets and Day 100 priorities for executive leadership",
            config, bold_prefix="Diligence-to-Integration Translation")
        add_bullet(doc,
            " — drove $4M+ organic growth through operational systems and patient re-engagement — "
            "filtered 10,000+ inactive records, triaged, tracked multi-channel outreach to visit conversion",
            config, bold_prefix="Revenue Growth & Value Realization")
        add_bullet(doc,
            " — managed technology integration across eClinicalWorks, Athenahealth, Curegram, Salesforce — "
            "consolidated 8 separate systems into unified operational platform",
            config, bold_prefix="Technology Integration")
        add_bullet(doc,
            " — identified execution risks early — governance rhythms and escalation paths that prevented "
            "operational friction from stalling deal momentum across every acquisition",
            config, bold_prefix="Execution Risk Management")
    elif company == "Brex":
        add_body(doc,
            "Led end-to-end operations scaling from 3→70 across 32 locations in 4 US states. "
            "Directed the infrastructure enabling $300K→$4M ARR and a $17M exit.", config, italic=True, size=Pt(9.5))
        add_bullet(doc,
            " — built decision infrastructure and scalable platforms from zero: hardware "
            "provisioned, VOIP/VPN/internet at 32 sites, EHR configured, Python data pipeline "
            "connecting EHR→billing→analytics — no prior playbook",
            config, bold_prefix="Built decision infrastructure from zero")
        add_bullet(doc,
            " — automated clinic scouting with Python — scraped real estate and provider data "
            "across 10+ states, replaced manual research with automated pipeline, accelerating "
            "target ID by 3x",
            config, bold_prefix="Automated clinic scouting")
        add_bullet(doc,
            " — deployed LLM/AI to overhaul communication — built prompt-based email refinement "
            "and automated patient-provider templates, cutting manual writing time by 60%",
            config, bold_prefix="Deployed AI-native automation")
        add_bullet(doc,
            " — scaled team 3→70 with sub-30-day new hire ramp, lead development, performance "
            "frameworks, zero-defect quality culture. Delivered quarterly investor-ready "
            "reporting to executive leadership.",
            config, bold_prefix="Scaled team 3→70")
        add_bullet(doc,
            " — directed 8 concurrent diligence workstreams through $17M acquisition — "
            "coordinated finance, legal, ops, IT, HR, compliance — clean close, 100% key "
            "talent retained through 18-month earnout",
            config, bold_prefix="Directed full-cycle acquisition")
        add_bullet(doc,
            " — consolidated 8 separate systems into one unified platform within 90 days — "
            "migrated EHR, billing, scheduling, reporting without service disruption",
            config, bold_prefix="Consolidated 8 systems")
        add_bullet(doc,
            " — presented quarterly board reports to 5-member executive committee — P&L "
            "variance across 12 departments, strategic tracking, financial packages governing "
            "$4M allocation. Advised CEO on acquisition strategy.",
            config, bold_prefix="Board-level strategic reporting")
    elif company == "Seaspan":
        add_body(doc,
            "Built organizational infrastructure during 23x hypergrowth — managing change across "
            "5 clinic groups, 12 departments, 32 locations through 5 acquisitions and a $17M exit.",
            config, italic=True, size=Pt(9.5))
        add_bullet(doc,
            " — managed organizational change through 5 acquisitions: built adoption frameworks, "
            "stakeholder communication plans, culture integration playbooks — 100% key talent retention",
            config, bold_prefix="Change Management & Organizational Transformation")
        add_bullet(doc,
            " — built change governance: weekly steering cadences, escalation paths, readiness "
            "checklists — every new location reached operational steady state within 30 days",
            config, bold_prefix="Change Governance & Readiness Frameworks")
        add_bullet(doc,
            " — stakeholder engagement across 12 departments with competing priorities: monthly "
            "check-ins, escalation frameworks, prioritization matrices — aligned 5 clinic groups",
            config, bold_prefix="Stakeholder Engagement & Alignment")
        add_bullet(doc,
            " — managed full P&L for $4M ARR: budget planning, variance analysis, multi-scenario "
            "forecasting — board-ready reporting to executive leadership",
            config, bold_prefix="P&L Management & Financial Reporting")
        add_bullet(doc,
            " — designed and delivered training for 70 FTEs across 32 locations — reduced new hire "
            "ramp from 60 to 30 days through structured onboarding and mentorship frameworks",
            config, bold_prefix="Training & Capability Building")
        add_bullet(doc,
            " — presented quarterly board reports to executive committee — P&L variance, strategic "
            "tracking, resource allocation governing $4M annual budget across 12 departments",
            config, bold_prefix="Executive Reporting & Board Communication")
    else:
        add_body(doc,
            "Directed end-to-end operations for a multi-site healthcare group. Served as the primary "
            "bridge between executive leadership and all operational teams.", config, italic=True, size=Pt(9.5))
        add_bullet(doc,
            " — 8 concurrent due diligence workstreams, integration playbook, Day 1/100 milestones, "
            "100% key talent retention through $17M transition",
            config, bold_prefix="Full-Cycle Acquisition Execution")

    # Earlier Career (add for DoorDash)
    if True:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run("Earlier Career")
        run.font.name = config["font"]
        run.font.size = config["size"]
        run.bold = True
        add_bullet(doc, "Digital Strategy Manager (2016–2018) — led digital strategy, campaign analytics, and ROI measurement; built reporting dashboards, optimized $500K+ annual ad spend", config)
        add_bullet(doc, "Client Services Representative (2014–2016) — managed client escalations for enterprise accounts; developed response protocols that reduced resolution time by 30%", config)
    # Additional Key Achievements for Clio (executive impact)
    add_section_header(doc, "Key Achievements", config)
    add_bullet(doc, "Scaled multi‑site operations from 3 to 70 employees across 32 locations, delivering a $17 M exit", config)
    add_bullet(doc, "Built end‑to‑end SaaS and cloud‑native infrastructure, data pipelines, and AI‑augmented workflows", config)
    add_bullet(doc, "Led $17 M acquisition with 8 diligence workstreams, 100 % talent retention, and a 90‑day integration plan", config)
    add_bullet(doc, "Implemented strategic financial modeling, P&L ownership, and capital allocation for a $4 M ARR business", config)
    add_bullet(doc, "Directed cross‑functional teams delivering 23× growth without adding complexity", config)

    # Education
    add_section_header(doc, "Education", config)
    if company == "DoorDash_Canada":
        add_body(doc, "MBA, Strategy & Finance  |  BSc, Information Technology  |  Post-Bacc Diploma, KPU", config, size=Pt(9), space_after=0)
    elif company == "Brex":
        add_body(doc, "MBA, International Business & IT  |  Post-Bacc Diploma, Technical Mgmt & Services — KPU  |  BSc, Information Technology", config, size=Pt(9), space_after=0)
    else:
        add_body(doc, "Master of Business Administration (MBA)  |  Post-Bacc Diploma, Technical Mgmt & Services — KPU  |  Post-Grad Diploma, Business Mgmt (IT)  |  BSc, Information Technology", config, size=Pt(9), space_after=0)

    # Technical Proficiency
    add_section_header(doc, "Technical Proficiency", config)
    if company == "Methanex":
        add_body(doc,
            "Financial Modelling & Analysis  ·  MS Excel (Advanced)  ·  Google Workspace  ·  "
            "OKR Frameworks  ·  ERP Systems  ·  CRM Platforms  ·  Jira/Confluence  ·  Data Visualization",
            config, size=Pt(9))
    elif company == "Hiive":
        add_body(doc,
            "Systems Architecture & Automation  ·  EHR/Practice Management  ·  "
            "Google Workspace  ·  KPI Dashboards  ·  Financial Modeling  ·  "
            "OKR Frameworks  ·  CRM  ·  Jira/Confluence  ·  AI-Augmented Workflows",
            config, size=Pt(9))
    elif company == "Providence_Healthcare":
        add_body(doc,
            "EHR/Practice Management  ·  Financial Modeling & Budgeting  ·  "
            "KPI Dashboards  ·  Google Workspace  ·  OKR Frameworks  ·  "
            "Project Management  ·  Regulatory Compliance  ·  Quality Improvement",
            config, size=Pt(9))
    elif company == "DoorDash_Canada":
        add_body(doc,
            "ERP/Financial Systems  ·  SQL  ·  Excel/Sheets (Advanced)  ·  "
            "Data Analysis  ·  Project Management",
            config, size=Pt(9))
    elif company == "Practice_Better":
        add_body(doc,
            "RevOps Platforms  ·  EHR (Athenahealth, eClinicalWorks)  ·  G Suite  ·  CRM  ·  "
            "KPI Dashboards  ·  Financial Modeling  ·  OKR Frameworks  ·  Jira/Confluence",
            config, size=Pt(9))
    elif company == "BWZ":
        add_body(doc,
            "Business Modeling (Excel/Sheets)  ·  SQL & Data Analytics  ·  KPI Dashboards  ·  "
            "AI-Augmented Workflows  ·  Project Management  ·  OKR Frameworks  ·  "
            "ERP/Financial Systems  ·  Data Visualization",
            config, size=Pt(9))
    elif company == "KPMG":
        add_body(doc,
            "Program Management & Governance  ·  ERP/Financial Systems  ·  Digital Transformation  ·  "
            "KPI Dashboards  ·  Financial Modeling  ·  AI-Augmented Workflows  ·  "
            "OKR Frameworks  ·  Jira/Confluence",
            config, size=Pt(9))
    elif company == "Microsoft":
        add_body(doc,
            "Program Management & Governance  ·  Operational Readiness  ·  "
            "KPI Dashboards  ·  AI-Augmented Workflows  ·  ERP/Financial  ·  "
            "OKR Frameworks  ·  Jira/Confluence  ·  Executive Communication",
            config, size=Pt(9))
    elif company == "Accenture":
        add_body(doc,
            "Strategic Transformation Frameworks  ·  Financial Modeling  ·  "
            "KPI Dashboards  ·  AI-Augmented Workflows  ·  ERP/Financial  ·  "
            "OKR Frameworks  ·  Jira/Confluence  ·  Board Reporting",
            config, size=Pt(9))
    elif company == "Brex":
        add_body(doc,
            "Python · SQL · LLM/AI (ChatGPT, Claude) · Salesforce · Retool · EHR (eClinicalWorks, Athenahealth) · "
            "VOIP/VPN · Cloud · Workflow Design · Cross-Functional Program Mgmt",
            config, size=Pt(9))
    elif company == "Seaspan":
        add_body(doc,
            "Change Management (ADKAR, Kotter, PROSCI-aligned)  ·  Project Management  ·  "
            "Jira/Confluence  ·  OKR Frameworks  ·  KPI Dashboards  ·  "
            "Google Workspace  ·  EHR/Operational Platforms  ·  Data Visualization",
            config, size=Pt(9))
    elif company == "Indeed":
        add_body(doc,
            "Athenahealth  ·  eClinicalWorks  ·  CRM  ·  Google Workspace  ·  "
            "Financial Modeling  ·  OKR Frameworks  ·  Jira/Confluence",
            config, size=Pt(9))

    if company == "Methanex":
        role_str = "Director_Strategy"
    elif company == "Hiive":
        role_str = "Associate_Operations_Strategy"
    elif company == "Providence_Healthcare":
        role_str = "Director_Clinical_Operations"
    elif company == "DoorDash_Canada":
        role_str = "Manager_SO_Dasher_Logistics"
    elif company == "UBC":
        role_str = "SrMgr_Strategic_Initiatives"
    elif company == "Practice_Better":
        role_str = "Director_Revenue_Operations"
    elif company == "BWZ":
        role_str = "Strategy_Ops_Manager"
    elif company == "KPMG":
        role_str = "Director_Delivery_Services"
    elif company == "Microsoft":
        role_str = "SrTPM_Security"
    elif company == "Accenture":
        role_str = "Performance_Strategy_Manager"
    elif company == "Brex":
        role_str = "BizOps_SrMgr_Technical"
    elif company == "Seaspan":
        role_str = "Change_Management_Specialist"
    else:
        role_str = "SrMgr_Integration"

    respath = os.path.join(folder, f"Aman_Kumar_{company}_{role_str}.docx")
    try:
        doc.save(respath)
    except PermissionError:
        respath = os.path.join(lfolder, os.path.basename(respath))
        doc.save(respath)
    doc.save(os.path.join(lfolder, os.path.basename(respath)))
    print(f"Resume: {respath}")

    # --- COVER LETTER ---
    doc = Document()
    set_margins(doc, config["margins"])

    add_body(doc, NAME, config, bold=True, space_after=0)
    cover_contact = doc.add_paragraph()
    cover_contact.paragraph_format.space_after = Pt(0)
    cover_contact.paragraph_format.space_before = Pt(0)
    add_plain_run(cover_contact, f"{PHONE}  |  ", config["font"], Pt(9), color="505050")
    add_hyperlink_contact(cover_contact, EMAIL, f"mailto:{EMAIL}", config["font"], Pt(9))
    add_plain_run(cover_contact, "  |  ", config["font"], Pt(9), color="505050")
    add_hyperlink_contact(cover_contact, "LinkedIn", f"https://{LINKEDIN}", config["font"], Pt(9))
    add_body(doc, LOCATION, config, size=Pt(9), space_after=4)
    DATE_LABELS = {"Indeed": "July 5, 2026", "Methanex": "June 21, 2026", "Deloitte": "June 19, 2026", "Hiive": "June 22, 2026", "Providence_Healthcare": "June 22, 2026", "DoorDash_Canada": "June 24, 2026", "UBC": "June 25, 2026", "BWZ": "June 25, 2026", "KPMG": "June 25, 2026", "Microsoft": "June 25, 2026", "Accenture": "July 7, 2026", "Brex": "July 15, 2026", "Seaspan": "July 14, 2026"}
    add_body(doc, DATE_LABELS.get(company, "June 22, 2026"), config, space_after=8)

    if company == "Methanex":
        add_body(doc, "Methanex Corporation", config, space_after=0)
        add_body(doc, "1800 Waterfront Centre, 200 Burrard Street", config, space_after=0)
        add_body(doc, "Vancouver, BC V6C 3M1", config, space_after=8)
        add_body(doc, "Re: Director, Strategy", config, bold=True, space_after=8)

        body = (
            "Dear Hiring Manager,\n\n"
            "Methanex is the world's largest methanol producer, navigating volatile commodity markets, "
            "the integration of a $2.05B acquisition, and a strategic shift toward low-carbon solutions. "
            "The Director, Strategy role sits at the center of these dynamics. This is exactly where I deliver the most value.\n\n"
            "I spent eight years building the strategic infrastructure of a multi-site organization from scratch. "
            "When I joined, there were 3 people and no playbook. When I left, it was 70 people across 32 locations "
            "with a $17M acquisition that I directed end-to-end. The strategic planning system I designed — annual "
            "strategy sessions cascading through quarterly OKRs with board-level reporting — gave executive leadership "
            "real-time visibility into execution across every location. The financial models I built governed resource "
            "allocation across 12 departments. The acquisition I directed involved 8 workstreams, diligence across "
            "finance, legal, and operations, and an integration that retained 100% of our key talent.\n\n"
            "What made this possible was not a pre-existing framework. I built it — the strategy cycle, the governance "
            "rhythms, the valuation models, the integration playbook — all from zero. I learned that strategy is not "
            "a document you produce once a year. It is a living process that must connect the Board room to the front "
            "line, and it requires someone who can operate at both altitudes without losing coherence.\n\n"
            "I understand Methanex operates in a different industry. But strategy is fractal — the core problems are "
            "the same: Where do we allocate capital? How do we grow? What risks do we manage? How do we align a global "
            "organization around a shared plan? I have solved these problems in my domain, and I am ready to solve them in yours.\n\n"
            "I would welcome the opportunity to discuss how my experience building strategic and financial infrastructure "
            "can support Methanex's next phase of global leadership."
        )
    elif company == "Providence_Healthcare":
        add_body(doc, "Providence Health Care", config, space_after=0)
        add_body(doc, "Burnaby, BC", config, space_after=8)
        add_body(doc, "Re: Director, Clinical and Operations (LTC PHC & FHA)", config, bold=True, space_after=8)

        body = (
            "Dear Hiring Committee,\n\n"
            "I spent 8 years leading multi-site healthcare operations — and the patient population I served? "
            "Seniors. I have done exactly what this role requires.\n\n"
            "At SkyflyMD, I served as the de facto Director of Operations for a multi-site healthcare organization "
            "that grew from 3 to 70 people across 32 locations. The practice was geriatrics-focused — I managed operations "
            "alongside a geriatrician, coordinating care across clinic locations and senior care homes, managing "
            "billing and compliance for aging populations, and building the systems that supported quality care "
            "for older adults. I managed the full P&L ($4M ARR). I built the operational infrastructure from "
            "scratch — scheduling, billing, compliance, quality assurance, reporting systems. I led interdisciplinary "
            "teams across 12 departments. I directed the $17M acquisition that integrated 8 separate operational "
            "systems into one without losing a single key team member.\n\n"
            "What drew me to Providence is the alignment between my experience and your mission. Multi-site seniors care "
            "at Chenchenstway and your PHC/FHA long-term care sites faces exactly the operational challenges I have been "
            "solving for the past 8 years: how to scale quality care across multiple locations serving aging populations "
            "without losing the personalized, compassionate approach that defines your organization.\n\n"
            "I hold an MBA with 8 years of progressive healthcare operations leadership — including direct experience "
            "serving geriatric populations in multi-facility settings. My combination of formal business education and "
            "hands-on seniors care operations experience provides the equivalent foundation this role requires.\n\n"
            "I would welcome the opportunity to discuss how my experience building and leading multi-site healthcare "
            "operations for aging populations can support Providence and Fraser Health's vision for seniors care."
        )
    elif company == "DoorDash_Canada":
        add_body(doc, "DoorDash Canada", config, space_after=0)
        add_body(doc, "Vancouver, BC / Toronto, ON", config, space_after=8)
        add_body(doc, "Re: Manager, Strategy & Operations – Dasher & Logistics Canada", config, bold=True, space_after=8)

        body = (
            "Dear Hiring Team,\n\n"
            "I built a business from nothing — scaled it to 70 people, 32 locations, "
            "managed every dollar of the P&L, and delivered a $17M exit. Along the way, "
            "I learned something that applies directly to this role: how you design pay "
            "and incentive systems determines whether your frontline workforce treats "
            "their work as a transaction or a partnership.\n\n"
            "At my company, I owned the complete compensation architecture for a "
            "distributed workforce across 32 locations. Every market had different "
            "dynamics — different cost of labour, different competitive pressure, "
            "different worker expectations. I built pay structures that balanced worker "
            "earnings with business cost efficiency, and I learned the behavioral "
            "economics of what actually motivates people in gig-like environments. "
            "Small changes in incentive design drove outsized shifts in performance.\n\n"
            "That experience maps directly to the challenge you're solving on the "
            "Dasher & Logistics team. DoorDash's marketplace depends on getting the "
            "pay and incentive equation right — attracting and retaining high-quality "
            "Dashers, at efficient spend, across diverse regional markets. It's a "
            "marketplace optimization problem that I've lived at operational scale.\n\n"
            "I'm drawn to DoorDash because your values match how I actually operate. "
            "\"Be an Owner\" — that's how I ran my business. \"Operate at the lowest "
            "level of detail\" — I still know the weekly cost per visit at every "
            "location I built. \"Bias for Action\" — I don't wait for perfect data "
            "to make a decision. I launch, test, iterate, and compound the wins.\n\n"
            "I'm ready to bring that owner-operator mindset to the Dasher & Logistics "
            "team. I understand the mechanics of labour supply, the tension between "
            "cost and quality, and the power of well-designed incentives.\n\n"
            "Let's talk."
        )
    elif company == "Hiive":
        add_body(doc, "Hiive", config, space_after=0)
        add_body(doc, "Vancouver, BC (HQ)", config, space_after=8)
        add_body(doc, "Re: Associate, Operations Strategy", config, bold=True, space_after=8)

        body = (
            "Dear Hiring Manager,\n\n"
            "I spent eight years doing exactly what this role describes: embedding with a revenue-generating team, "
            "identifying bottlenecks in the operational lifecycle, and deploying automated solutions to clear them.\n\n"
            "When I joined SkyflyMD, there was no operational infrastructure — just a team operating on willpower. "
            "By the time I directed the $17M exit, we had:\n\n"
            "- A complete tech stack (EHR, billing, scheduling, analytics) that I designed and implemented from scratch\n"
            "- Real-time KPI dashboards replacing manual spreadsheets\n"
            "- Automated workflows that reduced administrative overhead by 40%\n"
            "- A scalable operating system that supported 70 people across 32 locations — without adding operational complexity\n\n"
            "I did not inherit these systems. I built them. That is what systems builder means to me.\n\n"
            "Hiive is doing the same thing for the private market that I did for healthcare operations: replacing opacity "
            "with transparency, manual processes with automation, fragmentation with integration. The private secondary "
            "market has been running on brokers and spreadsheets for too long. You are building the infrastructure "
            "that changes that.\n\n"
            "I want to help you clear the bottlenecks.\n\n"
            "I am based in Vancouver and ready to be in your HQ five days a week."
        )
    elif company == "UBC":
        add_body(doc, "UBC Human Resources", config, space_after=0)
        add_body(doc, "University of British Columbia, Vancouver, BC", config, space_after=8)
        add_body(doc, "Re: Senior Manager, Strategic Initiatives and Engagement", config, bold=True, space_after=8)

        body = (
            "Dear Hiring Committee,\n\n"
            "UBC's Shaping UBC's Next Century plan articulates an ambitious mandate: deliver strategic priorities "
            "through disciplined resource allocation, stakeholder engagement, and institutional effectiveness. "
            "The Senior Manager, Strategic Initiatives and Engagement role sits at the center of this mandate. "
            "This is exactly where I deliver the most value.\n\n"
            "I spent eight years building the strategic and operational infrastructure for a multi-site organization "
            "from scratch. When I joined, there were 3 people, no governance structure, and no strategic planning "
            "process. When I left, it was 70 people across 32 locations with a $17M acquisition that I directed "
            "end-to-end. The strategic planning system I built --- annual strategy sessions cascading through "
            "departmental OKRs, board-level reporting, and quarterly performance reviews --- gave executive leadership "
            "real-time visibility into execution across every location. The budget frameworks I designed governed "
            "resource allocation across 12 departments. The acquisition I led involved 8 concurrent workstreams and "
            "an integration that retained 100% of key talent.\n\n"
            "What made this possible was not a pre-existing playbook. I created it --- the governance rhythms, "
            "the strategic planning cycles, the financial models, the stakeholder engagement frameworks, the policy "
            "documentation systems --- all from scratch, across a complex, multi-jurisdictional operating environment. "
            "I learned that strategic leadership in complex institutions is not about having all the answers. It is "
            "about building the process that surfaces the right questions, engages the right stakeholders, and delivers "
            "the right outcomes.\n\n"
            "UBC faces exactly the challenges I have solved: how to allocate constrained resources across competing "
            "strategic priorities, how to align diverse stakeholders around shared goals, how to build governance "
            "frameworks that satisfy both institutional accountability and external transparency, how to translate "
            "strategic intent into operational reality across a large, complex organization. I have navigated this "
            "complexity in healthcare. I am ready to navigate it at UBC.\n\n"
            "I am drawn to UBC because it is the intellectual backbone of British Columbia's future. I want to "
            "contribute to that future as someone who builds --- not as an observer, but as a practitioner who has "
            "delivered strategic infrastructure at scale. My MBA, my 8 years of progressive operational leadership, "
            "and my $17M acquisition experience have prepared me to do exactly what this role requires.\n\n"
            "I would welcome the opportunity to discuss how my experience building strategic and operational "
            "infrastructure can support UBC's next chapter of academic excellence and institutional impact."
        )
    elif company == "Practice_Better":
        add_body(doc, "Practice Better", config, space_after=0)
        add_body(doc, "Toronto, ON (Remote Canada)", config, space_after=8)
        add_body(doc, "Re: Director of Revenue Operations", config, bold=True, space_after=8)

        body = (
            "Dear Hiring Manager,\n\n"
            "I build revenue operations infrastructure. I've done it from zero, scaled it to $4M ARR, "
            "and delivered a $17M exit. That's the exact journey Practice Better is on right now --- "
            "and I want to be the operator who builds what comes next.\n\n"
            "When I joined SkyflyMD, there was no revenue operations system. No billing infrastructure, "
            "no scheduling automation, no KPI visibility. By the time I directed the $17M exit, we had "
            "a complete revenue operations stack spanning 32 locations: EHR, billing, RCM, analytics, "
            "and forecasting --- all built in-house, all scaled from zero to support 70 employees and "
            "$4M ARR. I didn't inherit a system. I built one.\n\n"
            "What makes this relevant to Practice Better is what I learned building that system: in healthcare "
            "SaaS, revenue operations is not just about pipeline management. It's about aligning clinical "
            "workflows with financial outcomes --- making sure the GTM engine and the product engine speak "
            "the same language. That's the gap I've seen most scaling healthcare SaaS companies struggle with, "
            "and that's exactly what I've solved.\n\n"
            "At SkyflyMD, I also managed the full P&L --- $4M ARR across 12 departments. I built the "
            "forecasting models, the variance analysis frameworks, and the board-level reporting that gave "
            "executive leadership real-time visibility. I know what it means to own the numbers and be "
            "accountable for them.\n\n"
            "I'm targeting Director of Revenue Operations roles at scaling healthcare SaaS companies because "
            "this is the exact problem I've already solved. Practice Better is at the stage where the founder "
            "can't own revenue ops alone anymore --- and that's where I come in. I build the function, hire "
            "the team, and create the infrastructure that makes the next phase of growth inevitable.\n\n"
            "I would welcome the opportunity to discuss how my experience building and scaling revenue "
            "operations for healthcare SaaS can support Practice Better's next chapter."
        )
    elif company == "KPMG":
        add_body(doc, "KPMG Canada", config, space_after=0)
        add_body(doc, "777 Dunsmuir Street, 11th Floor, Vancouver, BC V7Y 1K3", config, space_after=8)
        add_body(doc, "Re: Director, Delivery Services — Program Director (Req 32769)", config, bold=True, space_after=8)

        body = (
            "Dear Hiring Manager,\n\n"
            "KPMG's clients come to you to build the kind of business I've already built. "
            "I scaled a company from 3 people to 70, from zero to $4M ARR, from one location to 32 — "
            "and then directed a $17M acquisition that I managed end-to-end. "
            "I led the digital transformation, owned the P&L, built the governance frameworks, "
            "and retained 100% of key talent through the transition.\n\n"
            "What I've learned from that journey is directly relevant to the Director, Delivery Services role: "
            "program delivery is not about following a playbook — it's about building one when none exists. "
            "I designed the strategic planning system, the financial models, the governance rhythms, "
            "and the integration playbook from scratch — across 5 clinic groups, 12 departments, "
            "32 locations, and 5 jurisdictional regulatory environments.\n\n"
            "I'm drawn to KPMG because you serve clients who are navigating exactly these challenges — "
            "digital transformation, operational scaling, M&A integration, program governance. "
            "I've done what you advise. I want to bring that lived experience to your clients.\n\n"
            "I hold an MBA, a BSc in Information Technology, and 8 years of progressive leadership "
            "building and delivering complex programs at scale. I would welcome the opportunity to "
            "discuss how my program delivery experience can contribute to KPMG's Digital practice in Vancouver."
        )
    elif company == "BWZ":
        add_body(doc, "Black & White Zebra", config, space_after=0)
        add_body(doc, "Vancouver, BC (Remote)", config, space_after=8)
        add_body(doc, "Re: Strategy & Operations Manager", config, bold=True, space_after=8)

        body = (
            "Dear Hiring Manager,\n\n"
            "I read your JD and felt like someone had transcribed my career. 'Own the end-to-end build of zero-to-one initiatives.' "
            "'Bring order to the messy middle.' 'Impatient with meetings that don't move toward a decision.' "
            "That's not a job description — that's how I've operated for the last 8 years.\n\n"
            "I built a multi-site organization from 3 people to 70, across 32 locations, from zero to $4M ARR. "
            "I led the technology transformation, managed every dollar of the P&L, and directed a $17M exit. "
            "Along the way, I became the person people came to when they had a half-baked idea that needed "
            "pressure-testing — building the business case, running the costing, designing the go-to-market, "
            "and delivering the board-ready package that let leadership say yes or no with confidence.\n\n"
            "That's exactly what BWZ needs right now. You're at that sweet spot — 70 people, diversifying revenue, "
            "expanding beyond content into SaaS, and you need someone who can install project discipline "
            "without killing the startup soul. I've already walked this path. I know which processes matter "
            "and which ones just add bureaucracy.\n\n"
            "What draws me to BWZ specifically: you're bootstrapped and independent. You're not chasing unicorn "
            "valuations — you're building something sustainable. That's exactly how I built my company. "
            "And your CEO Ben Aston started as a solo blogger and built this from nothing — "
            "that's a founder story I deeply understand because I lived my own version of it.\n\n"
            "I'd welcome the chance to talk about how I can help BWZ bring order to the messy middle "
            "while keeping the builder spirit that got you here."
        )
    elif company == "Accenture":
        add_body(doc, "Accenture — Partners in Performance", config, space_after=0)
        add_body(doc, "Vancouver, BC", config, space_after=8)
        add_body(doc, "Re: Performance Strategy Manager — Req R00297351", config, bold=True, space_after=8)

        body = (
            "Dear Hiring Manager,\n\n"
            "Your clients come to you because they need to transform — and most of them have never done it before.\n"
            "I have. I've walked the path they're about to walk: building a business from nothing to something "
            "scalable, integrating acquisitions, leading through hypergrowth, and exiting with dignity and value.\n\n"
            "At SkyflyMD, I built the centralized operations backbone for a US medical practice roll-up that grew "
            "from 3 to 70 people across 32 locations in 4 states, integrated 5+ acquisitions, and delivered a "
            "$17M exit. I didn't advise from the sidelines — I was the one building the strategic planning "
            "system, designing the technology infrastructure, managing the P&L, and leading the team through "
            "every phase of transformation. I know what it feels like to be in the arena.\n\n"
            "That lived experience is what makes me different from a career consultant. When I sit with a client, "
            "I don't just recommend what they should do — I know exactly what it takes to implement it. I've "
            "navigated the messy middle between strategy and execution, and I can help your clients navigate it "
            "faster. I've already made the mistakes they're about to make, and I've built the playbooks that "
            "keep them from repeating them.\n\n"
            "What draws me to Partners in Performance specifically is your execution-first ethos. You don't just "
            "diagnose — you implement. That matches how I've operated my entire career. The 'embedded strategy "
            "partner' model is what I've been doing for 8 years, just inside one company instead of across many.\n\n"
            "I hold an MBA with deep P&L ownership, financial modeling, and board-level communication experience. "
            "I speak both the language of the C-suite and the language of the front line. I'm ready to bring that "
            "dual fluency to PiP's clients.\n\n"
            "I would welcome the opportunity to discuss how my experience building and transforming multi-site "
            "operations can bring immediate credibility and execution capability to Partners in Performance's "
            "engagements. I am based in Vancouver and willing to travel 80% as the role requires."
        )
    elif company == "Microsoft":
        add_body(doc, "Microsoft Security", config, space_after=0)
        add_body(doc, "Vancouver, BC", config, space_after=8)
        add_body(doc, "Re: Senior Technical Program Manager — Security", config, bold=True, space_after=8)

        body = (
            "Dear Hiring Manager,\n\n"
            "I read about Microsoft Security's 'Getting Customers Ready for AI' initiative and recognized "
            "something familiar: the challenge of defining operational readiness at scale for a complex, "
            "rapidly evolving domain. I've done exactly this before — not in security, but in healthcare, "
            "which has the same dynamics: regulation, distributed stakeholders, high-stakes outcomes, "
            "and the gap between wanting to adopt new technology and actually being ready to do it safely.\n\n"
            "I built a multi-site enterprise from 3 to 70 people, 32 locations, and $4M ARR — then "
            "directed a $17M exit. The defining achievement was not the growth itself but the operational "
            "infrastructure I built to make that growth sustainable: readiness frameworks that governed "
            "how we deployed technology, onboarded locations, trained teams, and measured success. "
            "Every new location followed the same process: assess readiness, build the baseline, deploy, "
            "measure, iterate. That framework turned chaos into repeatable scale.\n\n"
            "That's exactly what this role owns — defining the frameworks, metrics, and operational "
            "processes that help enterprises confidently prepare for AI adoption. The domains differ, "
            "but the skill is identical: building the system that makes readiness repeatable.\n\n"
            "What excites me about Microsoft specifically is the scope. Microsoft Security spans identity "
            "(Entra), data (Purview), threat protection (Defender), SIEM (Sentinel), and AI security "
            "(Security Copilot). Your Security Dashboard for AI (launched 2026) is exactly the kind of "
            "product that needs operational readiness frameworks to drive adoption at enterprise scale. "
            "The Forrester TEI study showing 124% ROI from security consolidation validates what I've "
            "always believed: readiness isn't a cost center — it's a value driver.\n\n"
            "I hold an MBA and 8 years building and scaling operational readiness in a complex, "
            "multi-stakeholder environment. I'm ready to bring that builder mindset to Microsoft Security.\n\n"
            "I would welcome the opportunity to discuss how my experience building readiness frameworks "
            "at scale can support Microsoft Security's mission to help enterprises adopt AI securely."
        )
    elif company == "Brex":
        add_body(doc, "Brex Talent Acquisition", config, space_after=0)
        add_body(doc, "Vancouver, BC", config, space_after=8)
        add_body(doc, "Re: BizOps Senior Manager (Technical)", config, bold=True, space_after=8)

        body = (
            "To the Brex Systems Team,\n\n"
            "I'm writing because your BizOps Senior Manager (Technical) role in Vancouver describes what I've "
            "spent the last eight years actually doing — building the systems, automation, and operational "
            "infrastructure that let an organization scale without breaking.\n\n"
            "When I joined SkyflyMD in 2018, it was three people in a single room. No SOPs. No infrastructure. "
            "No playbook for what came next. Over eight years, I walked that organization from zero to 70 people, "
            "32 locations, and a $17M exit — not by managing what existed, but by building what didn't. I designed "
            "the tech stack from scratch when there were no tools. I automated the workflows when manual couldn't "
            "scale. I directed the acquisition when the opportunity came.\n\n"
            "The part of your JD that resonates most: 'How should operational work be redesigned in an AI-native "
            "world?' I've been asking that since 2020 — writing Python scripts to automate clinic scouting years "
            "before the AI boom, deploying LLM tools that cut manual writing time by 60%, and teaching myself "
            "every system we needed because building means making the tools, not waiting for them.\n\n"
            "I know Brex is navigating the Capital One integration — which makes your Systems team's mandate "
            "even more critical. I've consolidated 8 separate operational systems into one platform in 90 days "
            "without disruption, retained 100% of key talent through a $17M acquisition transition, and built "
            "the integration playbook from scratch. That exact playbook is what I'd bring to the Systems team "
            "during this chapter.\n\n"
            "I don't describe myself as a manager who knows about systems. I'm a builder who leads through the "
            "systems he creates — which is why Brex's values — 'One Brex' (no silos), 'Dream Big' (think 10x), "
            "'Ownership' (if it's broken, you fix it) — describe how I've operated my entire career. I've scaled "
            "teams and written Python scripts in the same week because that's what operating at all levels means.\n\n"
            "I'd welcome the chance to walk your team through the systems I've built and how they apply "
            "to Brex's Operations challenges at scale."
        )
    else:
        add_body(doc, "Indeed", config, space_after=0)
        add_body(doc, "Vancouver, BC", config, space_after=8)
        add_body(doc, "Re: Sr. Manager, Integration & Business Acceleration — Reference ID: 47053", config, bold=True, space_after=8)

        body = (
            "Dear Hiring Manager,\n\n"
            "Your mission is to help people get jobs. Mine is to build the operational infrastructure "
            "that makes organizations scalable, efficient, and ready for their next phase of growth.\n\n"
            "I joined SkyflyMD when it was 3 people in a single location. When I left, it was 70 people "
            "across 32 locations, operating with systems I designed from scratch. The defining moment of "
            "that journey was directing our $17M acquisition — structuring 8 concurrent due diligence "
            "workstreams across finance, legal, operations, and provider contracts. I built the integration "
            "playbook, tracked Day 1 readiness and Day 100 milestones, consolidated 8 separate operational "
            "systems into one unified platform, and retained 100% of our key talent through the transition.\n\n"
            "What made that possible was not a pre-existing framework. There was no playbook. I created it — "
            "the governance rhythms, the reporting cadences, the escalation paths, the decision-making "
            "frameworks — all from scratch, across 5 clinic groups operating under different state regulations.\n\n"
            "What drew me to Indeed is that you are data-driven without being jargon-heavy, and mission-focused "
            "without being sentimental. That is how I operate. Your 'job seeker first' value resonates because "
            "it mirrors how I have always made operational decisions — starting with the end-user and working backward.\n\n"
            "I am not looking for a role that requires a playbook to exist before I start. I am looking for "
            "one that needs a playbook written.\n\n"
            "I would welcome the opportunity to discuss how my experience building and integrating multi-site "
            "operations can support Indeed's continued growth through M&A. Thank you for your time and consideration."
        )

    add_body(doc, body, config, space_after=0)
    add_body(doc, "", config, space_after=0)
    add_signature(doc, config)

    clpath = os.path.join(folder, f"Cover_Letter_{company}_{role_str}.docx")
    try:
        doc.save(clpath)
    except PermissionError:
        clpath = os.path.join(lfolder, os.path.basename(clpath))
        doc.save(clpath)
    doc.save(os.path.join(lfolder, os.path.basename(clpath)))
    print(f"Cover:   {clpath}")

if __name__ == "__main__":
    company = sys.argv[1] if len(sys.argv) > 1 else "Methanex"
    generate(company)
    # Final content validation — fail hard if output is too thin
    import glob as _gb
    out_dir = f"/mnt/c/Users/owner/OneDrive/ABHIMANYU-2.0/{get_date(company)}/{company}"
    docx_files = _gb.glob(os.path.join(out_dir, "Aman_Kumar_*.docx"))
    if docx_files:
        respath = docx_files[0]
        from docx import Document as DocCheck
        check = DocCheck(respath)
        c = count_docx_content(check)
        if c < 15:
            print(f"FATAL: Resume only has {c} meaningful lines (< 15 minimum). Content is too thin.")
            print(f"  → {'Company has content branch' if company in _COMPANIES_WITH_CONTENT else 'Company MISSING from _COMPANIES_WITH_CONTENT'}")
            print(f"  → {'SHOOT package found' if find_shoot_package(company) else 'No SHOOT package found'}")
            sys.exit(1)
        print(f"Content validated: {c} meaningful lines in resume.")
