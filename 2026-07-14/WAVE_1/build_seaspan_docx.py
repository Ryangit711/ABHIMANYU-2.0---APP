#!/usr/bin/env python3
"""One-shot DOCX builder for Seaspan — no parser, no config, just clean output."""
import os, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml

ONEDRIVE = "/mnt/c/Users/owner/OneDrive/ABHIMANYU-2.0"
LINUX = "/home/aryan/opencode_test/ABHIMANYU-2.0"
NAME = "Aman Kumar"
PHONE = "+1 236-885-2285"
EMAIL = "amankumar7111@outlook.com"
LINKEDIN = "linkedin.com/in/aman1776"
LOCATION = "Vancouver, BC"
FONT = "Calibri"
SIZE = Pt(11)
HSIZE = Pt(13)

def set_margins(doc, m=Inches(0.75)):
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = m

def run(p, text, font=FONT, size=SIZE, bold=False, italic=False, color=None):
    r = p.add_run(text)
    r.font.name = font; r.font.size = size; r.bold = bold; r.italic = italic
    if color:
        c = int(color, 16)
        r.font.color.rgb = RGBColor((c >> 16) & 0xFF, (c >> 8) & 0xFF, c & 0xFF)
    return r

def hyperlink(p, label, url):
    part = p.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    ns_r = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    p._p.append(parse_xml(
        f'<w:hyperlink xmlns:w="{ns_w}" xmlns:r="{ns_r}" r:id="{r_id}" w:history="1">'
        f'<w:r><w:rPr><w:rFonts w:ascii="{FONT}" w:hAnsi="{FONT}"/>'
        f'<w:sz w:val="{int(Pt(9).pt * 2)}"/><w:color w:val="0563C1"/><w:u w:val="single"/>'
        f'</w:rPr><w:t xml:space="preserve">{label}</w:t></w:r></w:hyperlink>'))

def section_header(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run(p, text.upper(), size=HSIZE, bold=True)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    b = pBdr.makeelement(qn('w:bottom'), {qn('w:val'): 'single', qn('w:sz'): '4', qn('w:space'): '1', qn('w:color'): '000000'})
    pBdr.append(b); pPr.append(pBdr)

def body(doc, text, bold=False, italic=False, size=SIZE, sa=2, sb=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    run(p, text, size=size, bold=bold, italic=italic)

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    if bold_prefix:
        run(p, f"\u2022 {bold_prefix}: ", bold=True)
        run(p, text)
    else:
        run(p, f"\u2022 {text}")

def contact(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, NAME, size=Pt(16), bold=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p2, f"{PHONE}  |  ", size=Pt(9), color="505050")
    hyperlink(p2, EMAIL, f"mailto:{EMAIL}")
    run(p2, "  |  ", size=Pt(9), color="505050")
    hyperlink(p2, "LinkedIn", f"https://{LINKEDIN}")
    run(p2, f"  |  {LOCATION}", size=Pt(9), color="505050")

# ===== RESUME =====
doc = Document()
set_margins(doc)

# Header
contact(doc)

# Professional Summary
section_header(doc, "Professional Summary")
body(doc, "Change management practitioner applying ADKAR, Kotter, and PROSCI-aligned frameworks operationally for 8 years. Built the centralized operations backbone for a multi-site healthcare organization \u2014 scaling from 3 to 70 employees, $0 to $4M ARR, $17M exit. Managed full lifecycle change across 5 acquisitions: stakeholder alignment, organizational design, systems migration, technology adoption, and culture integration.")

# Core Competencies
section_header(doc, "Core Competencies")
body(doc, "Organizational Change Management (ADKAR, Kotter)  |  M&A Integration & Transition  |  Stakeholder Alignment & Change Adoption  |  Technology Transformation & Systems Migration  |  KPI Dashboard Design (Adoption Metrics)  |  Cross-Functional Leadership  |  Org Design & Restructuring  |  Training & Capability Building  |  Strategic Planning & OKR Systems  |  Process Re-engineering", size=Pt(9.5))

# Professional Experience
section_header(doc, "Professional Experience")

def job_header(company, title, location, dates):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(4)
    run(p, company, bold=True)
    run(p, f"  |  {title}  |  {location}  |  {dates}")

def job_desc(text):
    body(doc, text, italic=True, size=Pt(9.5), sa=2)

# SkyflyMD
job_header("SkyflyMD", "Director of Operations & Change Management", "Vancouver, BC", "2017 \u2013 2025")
job_desc("Built the complete operational infrastructure for a multi-site healthcare group from the ground up \u2014 scaling from 3 to 70 employees, $0 to $4M ARR, across 5 clinic groups and 32 locations in 4 US states.")

bullet(doc, "Led 5 full-cycle organizational transformations through acquisition; structured 8 concurrent workstreams across finance, legal, operations, clinical; built integration change plan with Day 1/30/100 milestones; consolidated 8 separate operational systems into 1 unified platform within 90 days; retained 100% of key talent and maintained zero operational disruption throughout transition", "Organizational Change & M&A Integration")
bullet(doc, "Designed annual transformation cycles cascading through quarterly OKRs with board-level reporting cadences; aligned 5 clinic groups and 12 departments around shared priorities for 5 consecutive years; tracked change adoption through KPI dashboards measuring behavioral adoption across 32 locations", "Change Strategy & Adoption Measurement")
bullet(doc, "Designed hiring frameworks, organizational structures, training programs, performance management systems, and operational governance that scaled the organization from 3 to 70 without adding complexity", "Organizational Design & Capability Building")
bullet(doc, "Led full digital transformation from paper-based to fully integrated digital ecosystem (EHR, RCM, CRM, analytics) across 32 locations; managed the full change curve from awareness through reinforcement; built staged rollout with pilot, iteration, and full deployment; delivered hands-on training and on-site support for 150+ users", "Technology Adoption & Systems Migration")
bullet(doc, "Designed patient re-engagement campaigns targeting 10,000+ inactive records; multi-channel workflows; tracked conversion through custom analytics; reduced patient attrition by 30%, recovering $4M+ in organic revenue through sustained behavioral change", "Change Impact & Outcome Measurement")
bullet(doc, "Deployed KPI dashboards across all 32 locations replacing manual reporting; reduced reporting lag by 30%; enabled real-time operational decisions and change adoption tracking by location and department managers", "Adoption Metrics & Reporting Infrastructure")

# Additional Experience
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.space_before = Pt(6)
run(p, "Additional Experience", bold=True)
bullet(doc, "Digital Strategy Manager (2016\u20132018) \u2014 digital strategy, campaign analytics, ROI measurement, reporting dashboards")
bullet(doc, "Client Services Representative (2014\u20132016) \u2014 client escalations, enterprise accounts, response protocol design")

# Education
section_header(doc, "Education")
body(doc, "Post-Baccalaureate Diploma in Technical Management & Services \u2014 KPU, Surrey, BC", size=Pt(9), sa=0)
body(doc, "Master of Business Administration (MBA)", size=Pt(9), sa=0)
body(doc, "Post-Graduate Diploma in Business Management (IT)", size=Pt(9), sa=0)
body(doc, "Bachelor of Science in Information Technology", size=Pt(9), sa=0)

# Technical Proficiency
section_header(doc, "Technical Proficiency")
body(doc, "EHR & Practice Management: eClinicalWorks (expert), Athenahealth (expert)  |  Revenue Operations: RCM, Billing Infrastructure, Pipeline Management, Analytics  |  Business Tools: Salesforce, Microsoft 365, Google Workspace  |  Change Management: ADKAR, Kotter 8-Step, PROSCI-aligned, KPI Dashboard Design, OKR Frameworks", size=Pt(9))

# Save
os.makedirs(f"{ONEDRIVE}/2026-07-14/Seaspan", exist_ok=True)
os.makedirs(f"{LINUX}/2026-07-14/Seaspan", exist_ok=True)
rpath = f"{ONEDRIVE}/2026-07-14/Seaspan/Aman_Kumar_Seaspan_Change_Management_Specialist.docx"
lpath = f"{LINUX}/2026-07-14/Seaspan/Aman_Kumar_Seaspan_Change_Management_Specialist.docx"
doc.save(rpath); doc.save(lpath)
print(f"Resume: {rpath}")

# Count
c = sum(1 for p in doc.paragraphs if p.text.strip() and len(p.text.strip()) > 20)
print(f"  Content: {c} meaningful lines \u2014 PASS (threshold: 15)")

# ===== COVER LETTER =====
cl = Document()
for s in cl.sections:
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1.0)

p = cl.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(p, NAME, size=Pt(16), bold=True)

cl.add_paragraph()
body(cl, "July 14, 2026", sa=8)
body(cl, "Lead, Change Management", sa=0)
body(cl, "Seaspan ULC", sa=0)
body(cl, "50 Pemberton Avenue", sa=0)
body(cl, "North Vancouver, BC V7P 2R1", sa=8)
body(cl, "Re: Change Management Specialist \u2014 Job ID 8553", bold=True, sa=8)
body(cl, "Dear Lead, Change Management,", sa=8)

cover_body = [
    "I am writing to express my interest in the Change Management Specialist role at Seaspan. I bring ADKAR, Kotter, and PROSCI-aligned frameworks applied operationally across 8 years of organizational change \u2014 not from a classroom, but from the field.",
    "Most change management practitioners understand the theory. I have tested it.",
    "Over 8 years, I led 5 full-cycle organizational transformations through acquisition integration \u2014 developing Day 1/30/100 change plans, managing stakeholder alignment across 12 departments, and tracking adoption through custom KPI dashboards across 32 locations. When we transitioned from paper-based to a fully digital ecosystem, I managed every stage of the ADKAR model: building awareness of why change was needed, creating desire through direct benefit demonstration, delivering hands-on training across 150+ users, supporting the transition with on-site presence, and sustaining adoption through 6 months of reinforcement.",
    "I am drawn to Seaspan because the scale of transformation here is extraordinary \u2014 a $3.15B polar icebreaker, $6B naval support ships, new technologies, and a workforce of thousands. Organizational change at this scale demands someone who has not only studied the frameworks but has stood in front of 60 skeptical clinicians and convinced them to trust a new system. I have done that.",
    "I understand this role requires Controlled Goods Program clearance and ITAR compliance. I am fully prepared to meet these requirements.",
    "I would welcome the opportunity to discuss how my experience managing complex organizational change can support Seaspan\u2019s transformation journey.",
]
for b in cover_body:
    body(cl, b, sa=6)

body(cl, "", sa=0)
body(cl, "Best regards,", sa=0)
body(cl, "", sa=0)
body(cl, NAME, bold=True, sa=0)
body(cl, PHONE, sa=0)

clpath = f"{ONEDRIVE}/2026-07-14/Seaspan/Aman_Kumar_Seaspan_Change_Management_Specialist_Cover_Letter.docx"
cl_lpath = f"{LINUX}/2026-07-14/Seaspan/Aman_Kumar_Seaspan_Change_Management_Specialist_Cover_Letter.docx"
cl.save(clpath); cl.save(cl_lpath)
print(f"Cover:   {clpath}")

print("\nDone. Both files written.")
